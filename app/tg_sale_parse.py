"""Разбор реквизитов покупателя из PDF / Word / изображения."""
from __future__ import annotations

import base64
import json
import logging
import os
import re

log = logging.getLogger(__name__)

_FIELDS = (
    'name', 'inn', 'kpp', 'address', 'bank', 'rs', 'bik', 'ks',
)

_SYSTEM = (
    'Из текста или изображения реквизитов российской организации верни JSON: '
    '{"name":"","inn":"","kpp":"","address":"","bank":"","rs":"","bik":"","ks":""}. '
    'name — юрлицо или ФИО ИП. inn 10 или 12 цифр. kpp 9 цифр. '
    'rs — расчётный счёт 20 цифр. bik 9 цифр. ks — корсчёт. '
    'Пустая строка если поля нет. Без комментариев.'
)


def _digits(value, n=None) -> str:
    s = re.sub(r'\D+', '', str(value or ''))
    if n and len(s) != n:
        return s
    return s


def _clean(data: dict) -> dict:
    out = {k: str(data.get(k) or '').strip()[:500] for k in _FIELDS}
    out['inn'] = _digits(out['inn'])[:12]
    out['kpp'] = _digits(out['kpp'])[:9]
    out['rs'] = _digits(out['rs'])[:20]
    out['bik'] = _digits(out['bik'])[:9]
    out['ks'] = _digits(out['ks'])[:20]
    return out


def _extract_pdf_text(path: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        return ''
    chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages[:6]:
            chunks.append(page.extract_text() or '')
    return '\n'.join(chunks)


def _extract_docx_text(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        return ''
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            parts.append(' | '.join(c.text.strip() for c in row.cells))
    return '\n'.join(parts)


def _groq_from_text(text: str) -> dict:
    api_key = os.environ.get('GROQ_API_KEY', '').strip()
    if not api_key or not (text or '').strip():
        return {}
    try:
        from groq import Groq
        client = Groq(api_key=api_key)
        resp = client.chat.completions.create(
            model=os.environ.get('GROQ_MODEL', 'llama-3.3-70b-versatile'),
            messages=[
                {'role': 'system', 'content': _SYSTEM},
                {'role': 'user', 'content': text[:8000]},
            ],
            temperature=0.0,
            max_tokens=600,
            response_format={'type': 'json_object'},
        )
        return json.loads(resp.choices[0].message.content or '{}')
    except Exception:
        log.exception('tg_sale groq text parse failed')
        return {}


def _groq_from_image(data: bytes, mime: str) -> dict:
    api_key = os.environ.get('GROQ_API_KEY', '').strip()
    if not api_key or not data:
        return {}
    model = os.environ.get(
        'GROQ_VISION_MODEL',
        'meta-llama/llama-4-scout-17b-16e-instruct',
    )
    b64 = base64.b64encode(data).decode('ascii')
    url = f'data:{mime};base64,{b64}'
    try:
        from groq import Groq
        client = Groq(api_key=api_key)
        resp = client.chat.completions.create(
            model=model,
            messages=[{
                'role': 'user',
                'content': [
                    {'type': 'text', 'text': _SYSTEM},
                    {'type': 'image_url', 'image_url': {'url': url}},
                ],
            }],
            temperature=0.0,
            max_tokens=600,
            response_format={'type': 'json_object'},
        )
        return json.loads(resp.choices[0].message.content or '{}')
    except Exception:
        log.exception('tg_sale groq vision parse failed')
        return {}


def _heuristic(text: str) -> dict:
    body = (text or '').replace('\xa0', ' ')
    data = {k: '' for k in _FIELDS}
    m = re.search(r'ИНН\D{0,12}(\d{10,12})', body, re.I)
    if m:
        data['inn'] = m.group(1)
    m = re.search(r'КПП\D{0,12}(\d{9})', body, re.I)
    if m:
        data['kpp'] = m.group(1)
    m = re.search(r'БИК\D{0,12}(\d{9})', body, re.I)
    if m:
        data['bik'] = m.group(1)
    m = re.search(r'(?:р/?сч?|расч[её]тн)\D{0,24}(\d{20})', body, re.I)
    if m:
        data['rs'] = m.group(1)
    m = re.search(r'(?:к/?сч?|корр)\D{0,24}(\d{20})', body, re.I)
    if m:
        data['ks'] = m.group(1)
    m = re.search(
        r'(ООО|АО|ПАО|ИП)\s*[«"]?([^»"\n,]{3,80})',
        body,
    )
    if m:
        data['name'] = (m.group(0) or '').replace('«', '').replace('»', '').strip()[:200]
    return data


def parse_buyer_file(path: str, original_name: str = '', data: bytes | None = None) -> dict:
    """Возвращает {fields, error, source}."""
    name = (original_name or os.path.basename(path) or '').lower()
    ext = os.path.splitext(name)[1].lower()
    text = ''
    parsed = {}
    source = 'none'
    err = None

    if ext == '.pdf':
        text = _extract_pdf_text(path)
        source = 'pdf'
    elif ext in ('.doc', '.docx'):
        text = _extract_docx_text(path)
        source = 'docx'
        if not text:
            err = 'Не удалось прочитать Word. Впишите реквизиты вручную.'
    elif ext in ('.jpg', '.jpeg', '.png', '.webp', '.heic', '.bmp'):
        mime = {
            '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', '.png': 'image/png',
            '.webp': 'image/webp', '.bmp': 'image/bmp', '.heic': 'image/heic',
        }.get(ext, 'image/jpeg')
        blob = data
        if blob is None:
            with open(path, 'rb') as fh:
                blob = fh.read()
        parsed = _groq_from_image(blob, mime)
        source = 'image'
        if not any(parsed.values()):
            err = 'Фото не разобралось — проверьте поля вручную.'
    else:
        err = 'Нужен PDF, Word или фото реквизитов.'

    if text:
        parsed = _groq_from_text(text) or _heuristic(text)
        if not any(parsed.values()):
            parsed = _heuristic(text)
        if not parsed.get('name') and not parsed.get('inn'):
            err = err or 'Реквизиты не найдены — проверьте вручную.'

    fields = _clean(parsed)
    return {'fields': fields, 'error': err, 'source': source}
